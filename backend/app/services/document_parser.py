from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".json", ".png", ".jpg", ".jpeg"}

DOCUMENT_TYPES = {
    "invoice": "Invoice",
    "bank_statement": "Bank statement",
    "contract": "Contract",
    "receipt": "Receipt",
    "tax_document": "Tax document",
    "expense_report": "Expense report",
    "unknown": "Unknown document",
}

TYPE_KEYWORDS = {
    "invoice": ("invoice", "bill to", "invoice number", "amount due", "发票"),
    "bank_statement": ("bank statement", "account summary", "opening balance", "银行账单"),
    "contract": ("agreement", "contract", "terms and conditions", "合同"),
    "receipt": ("receipt", "payment received", "收据"),
    "tax_document": ("tax return", "tax form", "vat", "税务"),
    "expense_report": ("expense report", "reimbursement", "报销"),
}


@dataclass(slots=True)
class AnalysisResult:
    text: str
    document_type: str
    period: str | None
    confidence: float
    data: dict[str, Any]
    warnings: list[str]


def validate_extension(filename: str) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise ValueError(f"Unsupported file type. Allowed: {allowed}")
    return extension


def extract_text(content: bytes, extension: str) -> tuple[str, list[str]]:
    warnings: list[str] = []
    if extension == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        text = "\n\n".join(page for page in pages if page)
        if not text:
            warnings.append("No embedded PDF text found; OCR or manual review is required.")
        return text, warnings

    if extension == ".docx":
        document = DocxDocument(io.BytesIO(content))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        table_rows = [
            " | ".join(cell.text.strip() for cell in row.cells)
            for table in document.tables
            for row in table.rows
        ]
        return "\n".join(item for item in [*paragraphs, *table_rows] if item), warnings

    if extension == ".csv":
        decoded = _decode_text(content)
        rows = list(csv.reader(io.StringIO(decoded)))
        return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows[:250]), warnings

    if extension == ".json":
        decoded = _decode_text(content)
        try:
            parsed = json.loads(decoded)
            return json.dumps(parsed, ensure_ascii=False, indent=2)[:100_000], warnings
        except json.JSONDecodeError:
            warnings.append("The JSON file is invalid and requires manual review.")
            return decoded, warnings

    if extension in {".png", ".jpg", ".jpeg"}:
        warnings.append("Image OCR is not enabled in the local MVP; manual review is required.")
        return "", warnings

    return _decode_text(content), warnings


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def analyze_document(filename: str, content: bytes) -> AnalysisResult:
    extension = validate_extension(filename)
    text, warnings = extract_text(content, extension)
    searchable = f"{filename}\n{text}".lower()

    document_type, type_score = _classify(searchable)
    period = _extract_period(searchable)
    fields = _extract_fields(text, document_type)

    confidence = 0.1 if text else 0.0
    confidence += type_score
    confidence += 0.25 if period else 0.0
    confidence += min(0.2, len(fields) * 0.04)
    confidence = round(min(confidence, 0.98), 2)

    data: dict[str, Any] = {
        "document_type": document_type,
        "period": period,
        "fields": fields,
        "warnings": warnings,
        "text_characters": len(text),
        "extraction_method": "deterministic",
    }
    return AnalysisResult(
        text=text[:100_000],
        document_type=document_type,
        period=period,
        confidence=confidence,
        data=data,
        warnings=warnings,
    )


def _classify(searchable: str) -> tuple[str, float]:
    scores = {
        document_type: sum(1 for keyword in keywords if keyword in searchable)
        for document_type, keywords in TYPE_KEYWORDS.items()
    }
    winner, score = max(scores.items(), key=lambda item: item[1])
    if score == 0:
        return "unknown", 0.0
    return winner, min(0.45, 0.25 + score * 0.07)


def _extract_period(searchable: str) -> str | None:
    iso_match = re.search(r"\b(20\d{2})[-_/\.](0?[1-9]|1[0-2])\b", searchable)
    if iso_match:
        return f"{iso_match.group(1)}-{int(iso_match.group(2)):02d}"

    month_names = {
        "january": 1,
        "february": 2,
        "march": 3,
        "april": 4,
        "may": 5,
        "june": 6,
        "july": 7,
        "august": 8,
        "september": 9,
        "october": 10,
        "november": 11,
        "december": 12,
    }
    for name, month in month_names.items():
        match = re.search(rf"\b{name}\s+(20\d{{2}})\b", searchable)
        if match:
            return f"{match.group(1)}-{month:02d}"
    return None


def _extract_fields(text: str, document_type: str) -> dict[str, str]:
    fields: dict[str, str] = {}
    calendar_date = (
        r"((?:20\d{2}[-/][01]?\d[-/][0-3]?\d)|"
        r"(?:[A-Za-z]+\s+\d{1,2},?\s+20\d{2}))"
    )
    patterns = {
        "document_number": (
            r"(?:invoice|receipt|contract)\s*(?:number|no\.?|#)"
            r"\s*[:#-]?\s*([A-Z0-9-]{3,30})"
        ),
        "document_date": rf"(?:invoice date|date)\s*[:#-]?\s*{calendar_date}",
        "due_date": rf"due date\s*[:#-]?\s*{calendar_date}",
        "total_amount": (
            r"(?:total(?: amount)?|amount due|ending balance)"
            r"\s*[:#-]?\s*([€£$¥]?\s?[\d,]+(?:\.\d{2})?)"
        ),
        "account_last4": (
            r"(?:account|acct)(?:\s*(?:number|no\.?|#))?"
            r"\s*[:#-]?\s*(?:x+|\*+)?(\d{4})\b"
        ),
    }
    lowered = text.lower()
    for field, pattern in patterns.items():
        match = re.search(pattern, lowered, flags=re.IGNORECASE)
        if match:
            fields[field] = match.group(1).strip()

    currency_match = re.search(r"\b(USD|EUR|GBP|CNY|RMB|JPY)\b", text, flags=re.IGNORECASE)
    if currency_match:
        fields["currency"] = currency_match.group(1).upper()
    elif "$" in text:
        fields["currency"] = "USD"
    elif "€" in text:
        fields["currency"] = "EUR"
    elif "£" in text:
        fields["currency"] = "GBP"

    fields["classification"] = DOCUMENT_TYPES.get(document_type, document_type)
    return fields
